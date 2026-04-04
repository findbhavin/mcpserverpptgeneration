import os
import requests
import uuid
import tempfile
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_font(style, name="Aptos Narrow", size=None, bold=None, italic=None):
    font = style.font
    font.name = name
    if size:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    # For some fonts, Word requires setting the ascii theme font
    # but setting name directly usually works for modern Word versions
    
def _set_paragraph_format(style, line_spacing=None, space_before=None, space_after=None, alignment=None):
    pf = style.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if alignment is not None:
        pf.alignment = alignment

def apply_guidelines(doc_path: str, output_path: str):
    doc = Document(doc_path)

    # 11. Body Text
    # Font: Aptos Narrow, 12 pt. Line Spacing: 1.15 lines. Paragraph Spacing: 6 pt Before and 6 pt After
    normal_style = doc.styles['Normal']
    _set_font(normal_style, name="Aptos Narrow", size=12)
    _set_paragraph_format(normal_style, line_spacing=1.15, space_before=6, space_after=6)

    # 9. Heading Styles
    # Heading 1: 20 pt, Heading 2: 18 pt, Heading 3: 16 pt (all Aptos Narrow)
    for i, size in enumerate([20, 18, 16], start=1):
        heading_name = f'Heading {i}'
        if heading_name in doc.styles:
            h_style = doc.styles[heading_name]
            _set_font(h_style, name="Aptos Narrow", size=size)

    # 1. Table of Contents & 2. List of Tables & 3. List of Figures
    for style_name in ['TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'Table of Figures']:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            _set_font(style, name="Aptos Narrow", size=12, italic=(style_name == 'Table of Figures'))

    # Process all tables
    for table in doc.tables:
        # 5. Table Styles
        # Text Alignment: Left-aligned. Table Alignment: 0.5 indent from left margin.
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Python-docx doesn't easily support table indent directly through properties in some versions,
        # but we can try setting the left margin via tblPr
        tblPr = table._tbl.tblPr
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(Inches(0.5).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Borders: Black borders for all tables (We assume table style handles it, or apply manually)
        table.style = 'Table Grid'
        
        for row in table.rows:
            for cell in row.cells:
                # Cell Margins: Left / Right: 0.19 cm, Top / Bottom: 0.1 cm
                tcPr = cell._tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for margin, val in [('left', 0.19), ('right', 0.19), ('top', 0.1), ('bottom', 0.1)]:
                    node = OxmlElement(f'w:{margin}')
                    node.set(qn('w:w'), str(int(Cm(val).twips)))
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)
                
                # Font Inside Tables: Aptos Narrow, 10 pt. Text Alignment: Left-aligned.
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Aptos Narrow"
                        run.font.size = Pt(10)

    # Process paragraphs (Captions, Figures)
    for paragraph in doc.paragraphs:
        # Check if it's a Table Caption or Figure Caption
        # We can try to guess by checking style or content
        if paragraph.style.name == 'Caption':
            text_lower = paragraph.text.lower()
            if text_lower.startswith('table'):
                # 4. Table Captions
                # Left-aligned, 0.5 indent. Aptos Narrow, 12 pt, Regular. Spacing: 6 pt Before/After.
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Inches(0.5)
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)
                for run in paragraph.runs:
                    run.font.name = "Aptos Narrow"
                    run.font.size = Pt(12)
                    run.font.italic = False
            elif text_lower.startswith('figure') or text_lower.startswith('fig'):
                # 6. Figure Captions
                # Center-aligned. Aptos Narrow, 12 pt, Italic.
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Aptos Narrow"
                    run.font.size = Pt(12)
                    run.font.italic = True
                    
        # 7. Figures / Images
        # Center-aligned, Keep with next.
        # Check if paragraph contains an image (InlineShape)
        if 'Graphic' in paragraph._p.xml or 'pic:pic' in paragraph._p.xml:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True

    doc.save(output_path)
